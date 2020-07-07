import docx
from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import datetime
import re
import pandas as pd
import numpy as np
from glob import glob


def main():
    
    all_files_to_process = glob('word_documents/*.docx')
    print(all_files_to_process)
    #for doc in all_files_to_process:

    doc = docx.Document('word_documents\\2020 02 27 NR Daily Log.docx')

    docaslist = list()
    finallist = list()
    #doc = docx.Document(doc)
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            docaslist.append(block.text)
            
        #elif isinstance(block, Table):
        #    docaslist.append(table_print(block))
    
    docdf = cleanthelist(docaslist)
    dateofincident = getdate(doc)
    
    docdf = getrouteccil(docdf)
    #docdf = getlocation(docdf)
       
    docdf.insert(0,'incident_date' , dateofincident)
    exportfile(docdf,'interim_output\\','nrlog')


def getlocation(cleandoc):
    """
    This matches the locations in the daily log against a spreadsheet containing known locations and then adds geographical
    co-ordinates with duplicated data being removed and unknown locations being replaced by route names
    
    Parameters
    cleandoc:       A data frame holding the extract NR Log data

    Returns
    doc_with_geog:  A data frame holding the Log data with geographical data
    
    """

    #extract data from csv of locations, remove LUL data and sort data
    print("getting location information from 'location_data' folder")
    location_data = pd.read_csv('location_data\\location_data.csv', encoding='cp1252',usecols=['location_name','latitude','longitude','postcode','location_type'])
    location_data_no_LUL = location_data[location_data['location_type']!= 'LUL_station']
    location_list = list(location_data_no_LUL['location_name'])
    sorted_location_list = sorted(location_list)

    #look for place names after 'at','between' or 'approaching'
    print("looking for locations")
    location_candidates = []
    location_final = []
    for counter, incident in enumerate(cleandoc['narrative']):
        for loc in sorted_location_list:
            print(loc)
            if ('at ' + loc) in incident or ('between ' + loc) in incident or ('approaching ' + loc) in incident:
                location_candidates.append(loc)
        #remove duplicated location names
        distinct_locations = sorted(set(location_candidates),key = lambda x:location_candidates.index(x))
        
        #check for incidents where no location is found and replace with 'route' placeholder
        if not distinct_locations:
            final_location = 'route'
            
        else:
            #return the location name with the longest name
            final_location = max(distinct_locations,key=len)
        
        location_final.append(final_location)
        location_candidates = []
    
    #insert the location into the "found location" column
    cleandoc['found_location'] = location_final

    #merge the geographical data with daily log dataframe.  remove the unncessary location type column. drop duplicates
    doc_with_geog =  pd.merge(cleandoc,location_data_no_LUL,left_on='found_location',right_on='location_name',how='left')
    doc_with_geog.drop(['location_type'],axis=1,inplace=True)
    doc_with_geog = doc_with_geog.drop_duplicates()

    #replace the 'route' placeholder with the route column information
    doc_with_geog['found_location'] = np.where(doc_with_geog['found_location']=='route',doc_with_geog['route'],doc_with_geog['found_location'])

    return doc_with_geog



def getrouteccil(docdf):
    """
    This splits the data frame column to produce new columns holding route and ccil information

    Parameters
    docdf:      A dataframe holding the paragraphs of the document

    Returns
    docdf:      A dataframe with new columns in appropriate order

    """
    docdf[['route','narrative']] = docdf['narrative'].str.split(' – ',1,expand=True)
    docdf[['ccil','narrative']]  = docdf['narrative'].str.split(' / ',1,expand=True)
    
    docdf = docdf[['route','ccil','narrative']]

    print(docdf['ccil'])
    #replace blank ccils to be implemented
    
    

    return docdf
    

def getdate(docobj):
    """
    Gets date from the core properties of the document, with the first 10 characters being the date.
    converts string into date format

    Parameters:
    docobj:            A docx object containing the document

    Returns:
    dateofincident:    A datetimeobject representing the date of the incident
    """
    dateofincident = datetime.datetime.strptime(str(docobj.core_properties.title)[:10],'%Y %m %d').date()

    return dateofincident


def cleanthelist(text):
    """
    This takes the list of paragraphs from the word document and removes irrelevant entries.  It finds paragraphs with the key
    reference point CCIL and appends them and the following paragraph to a new list.  This new list is then converted to a dataframe
    
    Parameters
    text:       A docx Document object containing the full document

    Returns
    textdf:     A dataframe holding the relevant text documents
    """
    
    finallist = list()
    
    #remove non-reports
    cleanerdoc = list(filter(None,text))
    #remove first 26 items - the cover page
    #cleanerdoc = cleanerdoc[26:]

    cleanerdoc = [i for i in cleanerdoc if not i.startswith('None')]
    cleanerdoc = [i for i in cleanerdoc if not i.startswith('Disconnected')]

    #mask for the CCIL codes
    ccil = [i for i, s in enumerate(cleanerdoc) if 'CCIL' in s]

    #join ccil codes and ccil text
    for i in ccil:
        finallist.append(cleanerdoc[i] +" / "+ cleanerdoc[i+1])

    textdf = pd.DataFrame(finallist,columns=['narrative'])


    return textdf


##unashamedly stolen from https://github.com/python-openxml/python-docx/issues/276
def iter_block_items(parent):
    """
    Generate a reference to each paragraph and table child within *parent*,
    in document order. Each returned value is an instance of either Table or
    Paragraph. *parent* would most commonly be a reference to a main
    Document object, but also works for a _Cell object, which itself can
    contain paragraphs and tables.
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def table_print(block):
    tablelist = list()
    table=block
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                tablelist.append(paragraph.text)
               
    return tablelist


def exportfile(df,destinationpath,filename,numberoffiles=1):
    """
    This procedure exports the finalised file as a CSV file with a datetime stamp in filename

    Parameters:
    df        - a dataframe containing the finalised data
    destinationpath     - a string providing the filepath for the csv file
    numberoffiles       - an int with the number of files being processed
    
    Returns:
    None, but does export dataframe df as a csv object
    """
     
    formatted_date = datetime.datetime.now().strftime('%Y%m%d_%H-%M')
    destinationfilename = f'{filename}_{formatted_date}.csv'
    print(f"Exporting {filename} to {destinationpath}{destinationfilename}\n")
    print(f"If you want to check on progress, refresh the folder "+ destinationpath + " and check the size of the " + filename + ".csv file. \n")  
    df.to_csv(destinationpath + destinationfilename, encoding='cp1252',index=False)

if __name__ == '__main__':
    main()
